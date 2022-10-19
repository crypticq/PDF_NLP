import PyPDF2
import spacy
from spacy.lang.en.stop_words import STOP_WORDS
from string import punctuation
from heapq import nlargest
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import pdfplumber
from docx.shared import RGBColor


def get_pdf_text(pdf_file):
    text = []
    with pdfplumber.open(pdf_file) as pdf:
        for i in range(len(pdf.pages)):
            page = pdf.pages[i]
            text.append(page.extract_text())

    return text


def summarize(text, n=5):
    try:
        nlp = spacy.load('en_core_web_sm')
        doc = nlp(text)
        stopwords = list(STOP_WORDS)
        word_frequencies = {}
        for word in doc:
            if word.text not in stopwords:
                if word.text not in punctuation:
                    if word.text not in word_frequencies.keys():
                        word_frequencies[word.text] = 1
                    else:
                        word_frequencies[word.text] += 1

        max_frequency = max(word_frequencies.values())

        for word in word_frequencies.keys():
            word_frequencies[word] = (word_frequencies[word]/max_frequency)

        sentence_list = [sentence for sentence in doc.sents]

        sentence_scores = {}
        for sent in sentence_list:
            for word in sent:
                if word.text.lower() in word_frequencies.keys():
                    if len(sent.text.split(' ')) < 30:
                        if sent not in sentence_scores.keys():
                            sentence_scores[sent] = word_frequencies[word.text.lower()]
                        else:
                            sentence_scores[sent] += word_frequencies[word.text.lower()]

        select_length = int(len(sentence_list) * n)
        summary = nlargest(select_length, sentence_scores, key=sentence_scores.get)
        final_summary = [word.text for word in summary]
        summary = ' '.join(final_summary)
        return summary
    except Exception as e:
        print(e)
        pass

if __name__ == '__main__':
    res = []
    pdf_file = input('Enter the pdf file name: ')
    pdf_text = get_pdf_text(pdf_file)
    output_file = input('Enter the output file name without .docx: ')

    for text in pdf_text:
        data = (summarize(text))
        res.append(data)

    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.size = Pt(14)
    font.color.rgb = RGBColor.from_string('FF0000')
    document.add_heading('Summry', level=1)
    for w in final_summary:
        document.add_paragraph(w)

    document.save('{}.docx'.format(output_file))
    print('Summary is saved in {}.docx'.format(output_file))




