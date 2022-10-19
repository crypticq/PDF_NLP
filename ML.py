import pdfplumber
from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt

from transformers import BartTokenizer, BartForConditionalGeneration, BartConfig


def get_pdf_text(pdf_file):
    text = []
    with pdfplumber.open(pdf_file) as pdf:
        for i in range(len(pdf.pages)):
            page = pdf.pages[i]
            text.append(page.extract_text())

    return text

if __name__ == '__main__':

    pdf_file = input('Enter the pdf file name: ')
    pdf_text = get_pdf_text(pdf_file)
    output_file = input('Enter the output file name: ')
    print('Starting...')

    final_summary = []
    model = BartForConditionalGeneration.from_pretrained('sshleifer/distilbart-cnn-12-6')
    tokenizer = BartTokenizer.from_pretrained('sshleifer/distilbart-cnn-12-6')

    for text_in_pdf in pdf_text:
        inputs = tokenizer([text_in_pdf], truncation=True, return_tensors='pt')

        # Generate Summary
        summary_ids = model.generate(inputs['input_ids'], num_beams=4, early_stopping=True, min_length=0, max_length=1024)
        c = tokenizer.batch_decode(summary_ids, skip_special_tokens=True, clean_up_tokenization_spaces=False)[0]
        final_summary.append(c)

    print('Summary is ready')

    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.size = Pt(14)
    font.color.rgb = RGBColor.from_string('FF0000')
    document.add_heading('Summry', level=1)
    for w in final_summary:
        document.add_paragraph(w)

    document.save('{}.docx'.format(output_file))
    print('Summary is saved in {}.docx'.format(output_file))


